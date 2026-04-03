#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_diff_docx.py — Markdown Diff 转 Word 修订模式

将两个 Markdown 文件版本之间的差异转换为 Word 文档的修订（Track Changes）或批注（Comments）。

用法：
  python md_diff_docx.py old.md new.md -o diff.docx
  python md_diff_docx.py old.md new.md -o diff.docx --comments
  git diff HEAD~1 HEAD file.md | python md_diff_docx.py --from-diff -o diff.docx

依赖: pip install python-docx lxml latex2mathml
"""

import re
import os
import sys
import difflib
import unicodedata
from datetime import datetime, timezone
from typing import List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from md_core import (
    parse_md_blocks,
    HeadingBlock, ParagraphBlock, EquationBlock,
    TableBlock, CodeBlock, FigureBlock, ListBlock, BlankBlock,
    Block,
)


# ──────────────────────────────────────────────────────────
# 常量
# ──────────────────────────────────────────────────────────

COLOR_DELETE = 'FF0000'   # 红色
COLOR_INSERT = '00B050'   # 绿色
REVISION_AUTHOR = 'AutoDiff'
REVISION_DATE = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

_revision_id_counter = 0


def _next_revision_id() -> int:
    global _revision_id_counter
    _revision_id_counter += 1
    return _revision_id_counter


# ──────────────────────────────────────────────────────────
# Block 标准化（用于相似度比较）
# ──────────────────────────────────────────────────────────

def _block_key(block: Block) -> str:
    """将 Block 转换为用于比较的规范化字符串"""
    if isinstance(block, HeadingBlock):
        return f'H{block.level}:{block.text.strip()}'
    elif isinstance(block, ParagraphBlock):
        return block.text.strip()
    elif isinstance(block, EquationBlock):
        return f'EQ:{block.latex.strip()}'
    elif isinstance(block, TableBlock):
        return f'TABLE:{"|".join(block.header)}'
    elif isinstance(block, CodeBlock):
        return f'CODE:{block.code[:80].strip()}'
    elif isinstance(block, FigureBlock):
        return f'FIG:{block.path}'
    elif isinstance(block, ListBlock):
        return f'LIST:{"|".join(block.items)}'
    elif isinstance(block, BlankBlock):
        return ''
    return ''


def _block_text(block: Block) -> str:
    """提取 Block 的纯文本内容（用于词级 diff）"""
    if isinstance(block, HeadingBlock):
        return block.text
    elif isinstance(block, ParagraphBlock):
        return block.text
    elif isinstance(block, EquationBlock):
        return block.latex
    elif isinstance(block, CodeBlock):
        return block.code
    elif isinstance(block, ListBlock):
        return '\n'.join(block.items)
    return ''


def _has_formula(block: Block) -> bool:
    """检测 Block 是否包含公式（影响 track changes 可行性）"""
    if isinstance(block, EquationBlock):
        return True
    if isinstance(block, ParagraphBlock):
        return '$' in block.text
    return False


def _is_cjk_dominant(text: str) -> bool:
    """判断文本是否以中文为主（CJK 字符占比 > 50%）"""
    total = len(text)
    if total == 0:
        return False
    cjk_count = sum(
        1 for c in text
        if unicodedata.category(c) in ('Lo',) and ord(c) > 0x2E7F
    )
    return cjk_count / total > 0.5


# ──────────────────────────────────────────────────────────
# OOXML 工具函数
# ──────────────────────────────────────────────────────────

def _make_run_props(color: str, strikethrough: bool = False) -> OxmlElement:
    rpr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    rpr.append(color_el)
    if strikethrough:
        strike = OxmlElement('w:strike')
        rpr.append(strike)
    return rpr


def _make_w_t(text: str) -> OxmlElement:
    wt = OxmlElement('w:t')
    wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    wt.text = text
    return wt


def _make_del_text(text: str) -> OxmlElement:
    dt = OxmlElement('w:delText')
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    dt.text = text
    return dt


def _make_ins_element(text: str, author: str = REVISION_AUTHOR) -> OxmlElement:
    """创建 <w:ins> 元素（插入标记）"""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(_next_revision_id()))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), REVISION_DATE)
    r = OxmlElement('w:r')
    rpr = _make_run_props(COLOR_INSERT)
    r.append(rpr)
    r.append(_make_w_t(text))
    ins.append(r)
    return ins


def _make_del_element(text: str, author: str = REVISION_AUTHOR) -> OxmlElement:
    """创建 <w:del> 元素（删除标记）"""
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:id'), str(_next_revision_id()))
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), REVISION_DATE)
    r = OxmlElement('w:r')
    rpr = _make_run_props(COLOR_DELETE, strikethrough=True)
    r.append(rpr)
    r.append(_make_del_text(text))
    del_el.append(r)
    return del_el


def _make_normal_run(para, text: str):
    """向段落添加普通文本 run"""
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)


# ──────────────────────────────────────────────────────────
# 词/字符级 Diff
# ──────────────────────────────────────────────────────────

def _tokenize_text(text: str) -> List[str]:
    """将文本拆分为 token 序列（中文用字符级，英文用词级）"""
    if _is_cjk_dominant(text):
        return list(text)
    else:
        # 保留空格作为 token（确保 preserve 语义）
        tokens = re.split(r'(\s+)', text)
        return [t for t in tokens if t]


def _render_inline_diff(para, old_text: str, new_text: str, author: str):
    """在段落中渲染词级/字符级 diff（track changes 模式）"""
    old_tokens = _tokenize_text(old_text)
    new_tokens = _tokenize_text(new_text)

    sm = difflib.SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    for opcode, i1, i2, j1, j2 in sm.get_opcodes():
        if opcode == 'equal':
            _make_normal_run(para, ''.join(old_tokens[i1:i2]))
        elif opcode == 'delete':
            del_el = _make_del_element(''.join(old_tokens[i1:i2]), author)
            para._element.append(del_el)
        elif opcode == 'insert':
            ins_el = _make_ins_element(''.join(new_tokens[j1:j2]), author)
            para._element.append(ins_el)
        elif opcode == 'replace':
            del_el = _make_del_element(''.join(old_tokens[i1:i2]), author)
            para._element.append(del_el)
            ins_el = _make_ins_element(''.join(new_tokens[j1:j2]), author)
            para._element.append(ins_el)


# ──────────────────────────────────────────────────────────
# Unified Diff 解析（git diff 格式）
# ──────────────────────────────────────────────────────────

def _parse_unified_diff(diff_text: str) -> Tuple[List[str], List[str]]:
    """
    解析 unified diff 格式，还原出 old 和 new 的文本行列表。
    返回 (old_lines, new_lines)。
    """
    old_lines = []
    new_lines = []
    for line in diff_text.splitlines():
        if line.startswith('---') or line.startswith('+++'):
            continue
        if line.startswith('@@'):
            continue
        if line.startswith('-'):
            old_lines.append(line[1:])
        elif line.startswith('+'):
            new_lines.append(line[1:])
        else:
            # context line（空格开头）
            text = line[1:] if line.startswith(' ') else line
            old_lines.append(text)
            new_lines.append(text)
    return old_lines, new_lines


# ──────────────────────────────────────────────────────────
# Diff DOCX 渲染器
# ──────────────────────────────────────────────────────────

class DiffDocxRenderer:
    def __init__(self, use_comments: bool = False, author: str = REVISION_AUTHOR):
        self.doc = Document()
        self.use_comments = use_comments
        self.author = author
        self._setup_styles()

    def _setup_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        normal.font.size = Pt(10.5)
        normal.paragraph_format.line_spacing = 1.5

    # ── 渲染普通 Block ────────────────────────────────────

    def _render_normal_block(self, block: Block):
        """渲染未变化的 Block"""
        if isinstance(block, BlankBlock):
            return
        if isinstance(block, HeadingBlock):
            para = self.doc.add_heading(block.text, level=block.level)
            font_sizes = {1: 16, 2: 14, 3: 12, 4: 10.5, 5: 10.5}
            for run in para.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(font_sizes.get(block.level, 10.5))
                run.font.bold = True
        elif isinstance(block, ParagraphBlock):
            para = self.doc.add_paragraph()
            _make_normal_run(para, block.text)
        elif isinstance(block, EquationBlock):
            para = self.doc.add_paragraph()
            _make_normal_run(para, f'$${block.latex}$$')
        elif isinstance(block, TableBlock):
            # 简单渲染表格（不做三线表，保持可读性）
            if not block.header:
                return
            all_rows = [block.header] + block.rows
            table = self.doc.add_table(rows=len(all_rows), cols=len(block.header))
            table.style = 'Table Grid'
            for i, row_data in enumerate(all_rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_text
        elif isinstance(block, CodeBlock):
            para = self.doc.add_paragraph()
            run = para.add_run(block.code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif isinstance(block, FigureBlock):
            para = self.doc.add_paragraph()
            _make_normal_run(para, f'[图: {block.caption or block.alt or block.path}]')
        elif isinstance(block, ListBlock):
            for idx, item in enumerate(block.items):
                para = self.doc.add_paragraph()
                prefix = f'{idx + 1}. ' if block.ordered else '• '
                _make_normal_run(para, prefix + item)

    # ── Track Changes 渲染 ───────────────────────────────

    def _render_deleted_block(self, block: Block):
        """将整个 Block 标记为已删除（红色删除线）"""
        text = _block_text(block)
        if not text.strip():
            return
        label = self._block_label(block)
        para = self.doc.add_paragraph()
        if label:
            _make_normal_run(para, f'[{label}] ')
        del_el = _make_del_element(text, self.author)
        para._element.append(del_el)

    def _render_inserted_block(self, block: Block):
        """将整个 Block 标记为新增（绿色）"""
        text = _block_text(block)
        if not text.strip():
            return
        label = self._block_label(block)
        para = self.doc.add_paragraph()
        if label:
            _make_normal_run(para, f'[{label}] ')
        ins_el = _make_ins_element(text, self.author)
        para._element.append(ins_el)

    def _render_replaced_block(self, old_block: Block, new_block: Block):
        """渲染替换：对可文本化的 Block 做词级 diff，其他降级为批注"""
        # 表格、图片、公式块降级为 comment
        if (isinstance(old_block, (TableBlock, FigureBlock, EquationBlock)) or
                isinstance(new_block, (TableBlock, FigureBlock, EquationBlock))):
            self._render_replaced_as_comment(old_block, new_block)
            return

        # 含公式的段落降级为 comment
        if _has_formula(old_block) or _has_formula(new_block):
            self._render_replaced_as_comment(old_block, new_block)
            return

        old_text = _block_text(old_block)
        new_text = _block_text(new_block)

        # 标题级别变化
        if (isinstance(old_block, HeadingBlock) and isinstance(new_block, HeadingBlock)
                and old_block.level != new_block.level):
            para = self.doc.add_paragraph()
            _make_normal_run(para, f'[LEVEL CHANGE: H{old_block.level} → H{new_block.level}] ')
            _render_inline_diff(para, old_text, new_text, self.author)
            return

        # 普通段落词级 diff
        para = self.doc.add_paragraph()
        _render_inline_diff(para, old_text, new_text, self.author)

    def _render_replaced_as_comment(self, old_block: Block, new_block: Block):
        """用批注方式标记结构性变化"""
        # 渲染新内容为普通块
        self._render_normal_block(new_block)
        # 添加描述性段落
        old_text = _block_text(old_block)
        para = self.doc.add_paragraph()
        run = para.add_run(f'[CHANGED] 原内容: {old_text[:100]}{"..." if len(old_text) > 100 else ""}')
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)
        run.font.italic = True

    def _block_label(self, block: Block) -> str:
        if isinstance(block, HeadingBlock):
            return f'H{block.level}'
        if isinstance(block, TableBlock):
            return '表'
        if isinstance(block, CodeBlock):
            return '代码'
        if isinstance(block, FigureBlock):
            return '图'
        return ''

    # ── Comments 模式渲染 ─────────────────────────────────

    def _render_with_comment(self, block: Block, comment_text: str):
        """渲染块并附加批注（使用段落旁注方式，兼容 python-docx）"""
        self._render_normal_block(block)
        note_para = self.doc.add_paragraph()
        run = note_para.add_run(f'【批注】{comment_text}')
        run.font.color.rgb = RGBColor(0, 112, 192)
        run.font.size = Pt(9)
        run.font.italic = True

    # ── 主渲染入口 ────────────────────────────────────────

    def render_diff(self, old_blocks: List[Block], new_blocks: List[Block]):
        """渲染两组 Block 的差异"""
        # 过滤掉空行（空行在 diff 中意义不大）
        old_filtered = [b for b in old_blocks if not isinstance(b, BlankBlock)]
        new_filtered = [b for b in new_blocks if not isinstance(b, BlankBlock)]

        old_keys = [_block_key(b) for b in old_filtered]
        new_keys = [_block_key(b) for b in new_filtered]

        sm = difflib.SequenceMatcher(None, old_keys, new_keys, autojunk=False)
        opcodes = sm.get_opcodes()

        changed_count = sum(1 for op, *_ in opcodes if op != 'equal')
        if changed_count > 200:
            print(f"⚠  警告：检测到 {changed_count} 处变化，文档可能较大，Word 渲染可能变慢。")

        for opcode, i1, i2, j1, j2 in opcodes:
            if opcode == 'equal':
                for b in new_filtered[j1:j2]:
                    self._render_normal_block(b)

            elif opcode == 'delete':
                for b in old_filtered[i1:i2]:
                    if self.use_comments:
                        self._render_with_comment(b, f'[DELETED] 此段落已删除')
                    else:
                        self._render_deleted_block(b)

            elif opcode == 'insert':
                for b in new_filtered[j1:j2]:
                    if self.use_comments:
                        self._render_with_comment(b, f'[INSERTED] 此段落为新增内容')
                    else:
                        self._render_inserted_block(b)

            elif opcode == 'replace':
                old_chunk = old_filtered[i1:i2]
                new_chunk = new_filtered[j1:j2]

                if self.use_comments:
                    # 批注模式：渲染新内容 + 批注说明旧内容
                    for nb in new_chunk:
                        old_preview = '；'.join(
                            _block_text(ob)[:50] for ob in old_chunk
                        )
                        self._render_with_comment(
                            nb,
                            f'[CHANGED] 原内容: {old_preview}{"..." if len(old_preview) >= 50 else ""}'
                        )
                else:
                    # Track changes 模式
                    if len(old_chunk) == len(new_chunk) == 1:
                        # 单对单替换：词级 diff
                        ratio = difflib.SequenceMatcher(
                            None,
                            _block_key(old_chunk[0]),
                            _block_key(new_chunk[0])
                        ).ratio()
                        if ratio > 0.3:
                            self._render_replaced_block(old_chunk[0], new_chunk[0])
                        else:
                            self._render_deleted_block(old_chunk[0])
                            self._render_inserted_block(new_chunk[0])
                    else:
                        # 多对多：配对删除和插入
                        for b in old_chunk:
                            self._render_deleted_block(b)
                        for b in new_chunk:
                            self._render_inserted_block(b)

    def save(self, output_file: str):
        self.doc.save(output_file)
        print(f"✅ Diff 文档已保存: {output_file}")


# ──────────────────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────────────────

def diff_md_files(old_file: str, new_file: str, output_file: Optional[str] = None,
                  use_comments: bool = False, author: str = REVISION_AUTHOR) -> str:
    with open(old_file, 'r', encoding='utf-8') as f:
        old_text = f.read()
    with open(new_file, 'r', encoding='utf-8') as f:
        new_text = f.read()

    old_blocks = parse_md_blocks(old_text)
    new_blocks = parse_md_blocks(new_text)

    if output_file is None:
        base = os.path.splitext(new_file)[0]
        output_file = base + '_diff.docx'

    renderer = DiffDocxRenderer(use_comments=use_comments, author=author)
    renderer.render_diff(old_blocks, new_blocks)
    renderer.save(output_file)
    return output_file


def diff_from_unified(diff_text: str, output_file: str,
                      use_comments: bool = False, author: str = REVISION_AUTHOR) -> str:
    old_lines, new_lines = _parse_unified_diff(diff_text)
    old_blocks = parse_md_blocks('\n'.join(old_lines))
    new_blocks = parse_md_blocks('\n'.join(new_lines))

    renderer = DiffDocxRenderer(use_comments=use_comments, author=author)
    renderer.render_diff(old_blocks, new_blocks)
    renderer.save(output_file)
    return output_file


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(
        description='Markdown Diff → Word 修订文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python md_diff_docx.py old.md new.md -o diff.docx
  python md_diff_docx.py old.md new.md --comments -o diff.docx
  git diff HEAD~1 HEAD file.md | python md_diff_docx.py --from-diff -o diff.docx
        """
    )
    parser.add_argument('files', nargs='*', help='old.md new.md（两文件模式）')
    parser.add_argument('--from-diff', metavar='PATCH',
                        help='从 unified diff 文件读取（省略则从 stdin 读取）')
    parser.add_argument('-o', '--output', required=True, help='输出 DOCX 文件')
    parser.add_argument('--comments', action='store_true',
                        help='使用批注模式（默认为 Track Changes 修订模式）')
    parser.add_argument('--author', default=REVISION_AUTHOR,
                        help=f'修订作者名（默认: {REVISION_AUTHOR}）')
    args = parser.parse_args()

    if args.from_diff is not None:
        if args.from_diff == '-' or not args.from_diff:
            diff_text = sys.stdin.read()
        else:
            with open(args.from_diff, 'r', encoding='utf-8') as f:
                diff_text = f.read()
        diff_from_unified(diff_text, args.output, args.comments, args.author)
    elif len(args.files) == 2:
        diff_md_files(args.files[0], args.files[1], args.output,
                      args.comments, args.author)
    else:
        # 尝试从 stdin 读取 diff
        if not sys.stdin.isatty():
            diff_text = sys.stdin.read()
            diff_from_unified(diff_text, args.output, args.comments, args.author)
        else:
            parser.print_help()
            sys.exit(1)
