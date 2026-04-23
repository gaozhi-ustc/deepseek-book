#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 DOCX 工具 - LaTeX 公式自动转换版 v2
将 LaTeX 公式转换为 Word 原生公式（OMML）

依赖: pip install python-docx lxml latex2mathml
"""

import re
import os
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse
from lxml import etree
import latex2mathml.converter


MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# LaTeX 预处理
LATEX_SUBSTITUTIONS = [
    (r'\\text\{([^}]*)\}', r'\\mathrm{\1}'),
    (r'\\boldsymbol\{([^}]*)\}', r'\\mathbf{\1}'),
    (r'\\bm\{([^}]*)\}', r'\\mathbf{\1}'),
    (r'\\operatorname\{([^}]*)\}', r'\\mathrm{\1}'),
    (r'\\displaystyle', ''),
    (r'\\,', ' '),
    (r'\\;', ' '),
    (r'\\!', ''),
    (r'\\quad', '\\;'),
    (r'\\qquad', '\\;\\;'),
    (r'\\tag\{[^}]*\}', ''),
    (r'\\notag', ''),
    (r'\\label\{[^}]*\}', ''),
    (r'\\nonumber', ''),
]


def make_math_font_rpr():
    """创建公式字体属性 (Cambria Math)"""
    rpr = OxmlElement('m:rPr')
    sty = OxmlElement('m:sty')
    sty.set(qn('m:val'), 'p')  # plain
    rpr.append(sty)
    return rpr


def make_word_rpr_for_math():
    """创建 Word 运行属性，指定 Cambria Math 字体"""
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Cambria Math')
    rfonts.set(qn('w:hAnsi'), 'Cambria Math')
    rpr.append(rfonts)
    return rpr


def fix_omml_runs(element):
    """给 OMML 中的每个 m:r 添加字体属性"""
    for mr in element.iter(qn('m:r')):
        # 添加 w:rPr (Word run properties) 如果没有
        existing_wrpr = mr.find(qn('w:rPr'))
        if existing_wrpr is None:
            wrpr = make_word_rpr_for_math()
            # 插入到 m:t 之前
            mt = mr.find(qn('m:t'))
            if mt is not None:
                mr.insert(list(mr).index(mt), wrpr)
            else:
                mr.insert(0, wrpr)


class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.load_xsl()
        self.stats = {'success': 0, 'failed': 0, 'total': 0}
        
    def setup_styles(self):
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].paragraph_format.line_spacing = 1.5
    
    def load_xsl(self):
        xsl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MML2OMML.XSL')
        if not os.path.exists(xsl_path):
            raise FileNotFoundError(f"找不到 MML2OMML.XSL: {xsl_path}")
        xslt_doc = etree.parse(xsl_path)
        self.xslt_transform = etree.XSLT(xslt_doc)
    
    def preprocess_latex(self, latex_str):
        for pattern, replacement in LATEX_SUBSTITUTIONS:
            latex_str = re.sub(pattern, replacement, latex_str)
        latex_str = re.sub(r'\\begin\{(aligned|align\*?|gather\*?|equation\*?)\}', '', latex_str)
        latex_str = re.sub(r'\\end\{(aligned|align\*?|gather\*?|equation\*?)\}', '', latex_str)
        latex_str = latex_str.replace('&', ' ')
        latex_str = latex_str.replace('\\\\', ' ')
        return latex_str.strip()
    
    def latex_to_omml(self, latex_str):
        """LaTeX → MathML → OMML，返回可直接插入 docx 的元素"""
        self.stats['total'] += 1
        try:
            latex_str = latex_str.strip()
            if latex_str.startswith('$$') and latex_str.endswith('$$'):
                latex_str = latex_str[2:-2].strip()
            elif latex_str.startswith('$') and latex_str.endswith('$'):
                latex_str = latex_str[1:-1].strip()
            
            latex_str = self.preprocess_latex(latex_str)
            if not latex_str:
                return None
            
            # LaTeX → MathML
            mathml_str = latex2mathml.converter.convert(latex_str)
            
            # MathML → OMML
            mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
            omml_tree = self.xslt_transform(mathml_tree)
            omml_root = omml_tree.getroot()
            
            # 关键：将 lxml 元素转为 python-docx 兼容的元素
            omml_xml = etree.tostring(omml_root)
            omml_element = etree.fromstring(omml_xml)
            
            # 给每个 m:r 添加字体属性
            fix_omml_runs(omml_element)
            
            self.stats['success'] += 1
            return omml_element
        except Exception as e:
            self.stats['failed'] += 1
            preview = latex_str[:60] if len(latex_str) > 60 else latex_str
            print(f"  ⚠ 公式转换失败: {preview}  错误: {e}")
            return None
    
    def set_font(self, run, font_name='宋体', font_size=10.5, bold=False, color=None, italic=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    
    def add_heading(self, text, level):
        heading = self.doc.add_heading(text, level=level)
        font_sizes = {1: 16, 2: 14, 3: 12}
        for run in heading.runs:
            self.set_font(run, '黑体', font_sizes.get(level, 10.5), bold=True)
    
    def add_paragraph(self, text):
        """添加段落，行内公式转为 OMML"""
        if not text.strip():
            return
        
        para = self.doc.add_paragraph()
        parts = re.split(r'(\$[^$]+\$|`[^`]+`|\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('$') and part.endswith('$'):
                omml = self.latex_to_omml(part)
                if omml is not None:
                    para._element.append(omml)
                else:
                    run = para.add_run(part)
                    self.set_font(run, 'Cambria Math', 11, italic=True, color=RGBColor(255, 0, 0))
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                self.set_font(run, 'Courier New', 9)
            elif part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                self.set_font(run, '黑体', 10.5, bold=True)
            else:
                run = para.add_run(part)
                self.set_font(run)
    
    def add_equation(self, latex):
        """添加块级公式，用 m:oMathPara 包裹实现居中"""
        omml = self.latex_to_omml(latex)
        if omml is not None:
            # 创建 oMathPara（块级公式容器）
            omathpara = OxmlElement('m:oMathPara')
            omathpara_pr = OxmlElement('m:oMathParaPr')
            jc = OxmlElement('m:jc')
            jc.set(qn('m:val'), 'center')
            omathpara_pr.append(jc)
            omathpara.append(omathpara_pr)
            omathpara.append(omml)
            
            # 创建段落并插入
            para = self.doc.add_paragraph()
            para._element.append(omathpara)
        else:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(latex)
            self.set_font(run, 'Cambria Math', 11, italic=True, color=RGBColor(255, 0, 0))
    
    def add_table_from_markdown(self, lines):
        rows = []
        for line in lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if not all(re.match(r'^[-:]+$', cell) for cell in cells):
                    rows.append(cells)
        if len(rows) < 2:
            return
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j >= len(table.rows[i].cells):
                    continue
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_font(run, '宋体', 9)
                        if i == 0:
                            run.font.bold = True
    
    def add_code_block(self, code):
        para = self.doc.add_paragraph()
        run = para.add_run(code)
        self.set_font(run, 'Courier New', 9)
        para.paragraph_format.left_indent = Inches(0.5)
    
    def parse_markdown(self, md_file):
        print(f"正在转换: {md_file}")
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('```'):
                if in_code_block:
                    self.add_code_block('\n'.join(code_lines))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            if '|' in line and not in_table:
                in_table = True
                table_lines = [line]
                i += 1
                continue
            
            if in_table:
                if '|' in line:
                    table_lines.append(line)
                    i += 1
                    continue
                else:
                    self.add_table_from_markdown(table_lines)
                    table_lines = []
                    in_table = False
            
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                text = line.lstrip('#').strip()
                self.add_heading(text, level)
            elif line.strip().startswith('$$'):
                equation = line.strip()
                if not line.strip().endswith('$$') or line.strip() == '$$':
                    i += 1
                    eq_lines = []
                    while i < len(lines) and not lines[i].strip().endswith('$$'):
                        eq_lines.append(lines[i].rstrip())
                        i += 1
                    if i < len(lines):
                        eq_lines.append(lines[i].rstrip())
                    equation = equation + '\n' + '\n'.join(eq_lines)
                self.add_equation(equation)
            elif line.strip():
                self.add_paragraph(line)
            
            i += 1
        
        if in_table and table_lines:
            self.add_table_from_markdown(table_lines)
    
    def save(self, output_file):
        self.doc.save(output_file)
        s = self.stats
        print(f"\n✅ 文档已保存: {output_file}")
        print(f"📊 公式统计: 共 {s['total']} 个, 成功 {s['success']} ✓, 失败 {s['failed']} ✗")
        if s['failed'] > 0:
            print("💡 失败的公式以红色斜体显示在文档中")


def main():
    parser = argparse.ArgumentParser(description='Markdown → DOCX（LaTeX 公式 → Word 原生公式）')
    parser.add_argument('input', help='输入 Markdown 文件')
    parser.add_argument('-o', '--output', help='输出 DOCX 文件')
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 {args.input}")
        return
    
    output_file = args.output or os.path.splitext(args.input)[0] + '_formula.docx'
    
    converter = MarkdownToDocx()
    converter.parse_markdown(args.input)
    converter.save(output_file)


if __name__ == '__main__':
    main()
