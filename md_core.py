#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_core.py — Markdown 转 DOCX 核心转换器

两阶段架构：
  1. Tokenize：将 Markdown 解析为 Block 列表
  2. Render：将 Block 列表渲染为 DOCX

依赖: pip install python-docx lxml latex2mathml
"""

import re
import os
from dataclasses import dataclass, field
from typing import List, Optional, Literal, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter


# ──────────────────────────────────────────────────────────
# Block 数据模型
# ──────────────────────────────────────────────────────────

@dataclass
class Revision:
    kind: Literal['ins', 'del']
    text: str
    author: str
    date: str          # ISO 8601
    rev_id: int


@dataclass
class Comment:
    comment_id: int
    author: str
    date: str
    text: str                        # 批注正文
    anchor_text: str                 # 批注锚点原文
    anchor_range: Tuple[int, int]    # Block text 内字符偏移 [start, end)


@dataclass
class HeadingBlock:
    level: int
    text: str
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class ParagraphBlock:
    text: str
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class EquationBlock:
    latex: str       # 不含 $$ 分隔符的原始 LaTeX
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class TableBlock:
    header: List[str]
    rows: List[List[str]]
    caption: str     # 可能为空，由规范层填充
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class CodeBlock:
    code: str
    language: str
    title: str       # 可能为空，由规范层填充
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class FigureBlock:
    alt: str
    path: str
    caption: str     # 可能为空
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class ListBlock:
    items: List[str]
    ordered: bool
    raw: str
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

@dataclass
class BlankBlock:
    raw: str = ''
    revisions: List[Revision] = field(default_factory=list)
    comments: List[Comment]   = field(default_factory=list)

Block = (HeadingBlock | ParagraphBlock | EquationBlock | TableBlock |
         CodeBlock | FigureBlock | ListBlock | BlankBlock)


# ──────────────────────────────────────────────────────────
# LaTeX 预处理与转换
# ──────────────────────────────────────────────────────────

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

LATEX_SUBSTITUTIONS = [
    (r'\\text\{([^}]*)\}',        r'\\mathrm{\1}'),
    (r'\\boldsymbol\{([^}]*)\}',  r'\\mathbf{\1}'),
    (r'\\bm\{([^}]*)\}',          r'\\mathbf{\1}'),
    (r'\\operatorname\{([^}]*)\}',r'\\mathrm{\1}'),
    (r'\\displaystyle',           ''),
    (r'\\,',                      ' '),
    (r'\\;',                      ' '),
    (r'\\!',                      ''),
    (r'\\quad',                   '\\;'),
    (r'\\qquad',                  '\\;\\;'),
    (r'\\tag\{[^}]*\}',           ''),
    (r'\\notag',                  ''),
    (r'\\label\{[^}]*\}',         ''),
    (r'\\nonumber',               ''),
]


def preprocess_latex(latex_str: str) -> str:
    for pattern, replacement in LATEX_SUBSTITUTIONS:
        latex_str = re.sub(pattern, replacement, latex_str)
    return latex_str.strip()


def fix_omml_runs(element):
    """给 OMML 中的每个 m:r 添加 Cambria Math 字体属性"""
    for mr in element.iter(qn('m:r')):
        if mr.find(qn('w:rPr')) is None:
            rpr = OxmlElement('w:rPr')
            rfonts = OxmlElement('w:rFonts')
            rfonts.set(qn('w:ascii'), 'Cambria Math')
            rfonts.set(qn('w:hAnsi'), 'Cambria Math')
            rpr.append(rfonts)
            mt = mr.find(qn('m:t'))
            if mt is not None:
                mr.insert(list(mr).index(mt), rpr)
            else:
                mr.insert(0, rpr)


# ──────────────────────────────────────────────────────────
# Markdown Tokenizer
# ──────────────────────────────────────────────────────────

def _is_table_separator(line: str) -> bool:
    """判断是否是 Markdown 表格分隔行（如 |---|---|）"""
    stripped = line.strip()
    if not stripped.startswith('|'):
        return False
    cells = [c.strip() for c in stripped.split('|')[1:-1]]
    return len(cells) > 0 and all(re.match(r'^[-:]+$', c) for c in cells if c)


def _parse_table_row(line: str) -> List[str]:
    cells = [c.strip() for c in line.strip().split('|')]
    # 去掉首尾空字符串（由首尾 | 产生）
    if cells and cells[0] == '':
        cells = cells[1:]
    if cells and cells[-1] == '':
        cells = cells[:-1]
    return cells


def tokenize(md_text: str) -> List[Block]:
    """将 Markdown 文本解析为 Block 列表"""
    lines = md_text.splitlines()
    blocks: List[Block] = []
    i = 0

    # 用于检测前一个非空块是否是可能的表题/代码标题
    pending_caption: Optional[str] = None
    pending_code_title: Optional[str] = None

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ── 代码块 ─────────────────────────────────────────
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            raw_lines = [line]
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                raw_lines.append(lines[i])
                i += 1
            if i < len(lines):
                raw_lines.append(lines[i])
                i += 1
            title = pending_code_title or ''
            pending_code_title = None
            pending_caption = None
            blocks.append(CodeBlock(
                code='\n'.join(code_lines),
                language=lang,
                title=title,
                raw='\n'.join(raw_lines),
            ))
            continue

        # ── 块级公式 $$...$$ ───────────────────────────────
        if stripped.startswith('$$'):
            raw_lines = [line]
            # 单行 $$...$$ 且不是仅 $$
            if stripped.endswith('$$') and len(stripped) > 2:
                latex = stripped[2:-2].strip()
                i += 1
            else:
                # 收集开头 $$ 行 $$ 之后的内容（如 $$\begin{cases}）
                first_content = stripped[2:].strip()
                i += 1
                eq_lines = [first_content] if first_content else []
                while i < len(lines) and not lines[i].rstrip().endswith('$$'):
                    eq_lines.append(lines[i])
                    raw_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    raw_lines.append(lines[i])
                    last = lines[i].rstrip()
                    if last.endswith('$$'):
                        tail = last[:-2].strip()
                        if tail:
                            eq_lines.append(tail)
                    i += 1
                latex = '\n'.join(eq_lines)
            pending_caption = None
            pending_code_title = None
            blocks.append(EquationBlock(latex=latex.strip(), raw='\n'.join(raw_lines)))
            continue

        # ── 图片 ![alt](path) ──────────────────────────────
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]*)\)\s*$', stripped)
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2)
            raw_lines = [line]
            # 检查下一行是否是图题
            caption = ''
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and re.match(r'^图\s*\d+', lines[j].strip()):
                caption = lines[j].strip()
                raw_lines.append(lines[j])
                i = j + 1
            else:
                i += 1
            pending_caption = None
            pending_code_title = None
            blocks.append(FigureBlock(alt=alt, path=path, caption=caption, raw='\n'.join(raw_lines)))
            continue

        # ── 表格 ───────────────────────────────────────────
        if '|' in stripped and stripped.strip().startswith('|'):
            # 确认下一行是分隔行
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and _is_table_separator(lines[j]):
                raw_lines = []
                table_lines = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                    raw_lines.append(lines[i])
                    table_lines.append(lines[i])
                    i += 1
                # 解析表头（第0行）和分隔行（第1行）及数据行
                header = _parse_table_row(table_lines[0]) if table_lines else []
                rows = []
                for tl in table_lines[2:]:  # 跳过分隔行
                    rows.append(_parse_table_row(tl))
                caption = pending_caption or ''
                pending_caption = None
                pending_code_title = None
                blocks.append(TableBlock(
                    header=header,
                    rows=rows,
                    caption=caption,
                    raw='\n'.join(raw_lines),
                ))
                continue

        # ── 标题 # ─────────────────────────────────────────
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            pending_caption = None
            pending_code_title = None
            blocks.append(HeadingBlock(level=level, text=text, raw=line))
            i += 1
            continue

        # ── 无序列表 ───────────────────────────────────────
        if re.match(r'^[-*+]\s+', stripped):
            raw_lines = [line]
            items = [re.sub(r'^[-*+]\s+', '', stripped)]
            i += 1
            while i < len(lines) and re.match(r'^[-*+]\s+', lines[i].rstrip()):
                raw_lines.append(lines[i])
                items.append(re.sub(r'^[-*+]\s+', '', lines[i].rstrip()))
                i += 1
            pending_caption = None
            pending_code_title = None
            blocks.append(ListBlock(items=items, ordered=False, raw='\n'.join(raw_lines)))
            continue

        # ── 有序列表 ───────────────────────────────────────
        if re.match(r'^\d+\.\s+', stripped):
            # 先检测是否像 4级标题（规范层会进一步处理）
            raw_lines = [line]
            items = [re.sub(r'^\d+\.\s+', '', stripped)]
            i += 1
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].rstrip()):
                raw_lines.append(lines[i])
                items.append(re.sub(r'^\d+\.\s+', '', lines[i].rstrip()))
                i += 1
            pending_caption = None
            pending_code_title = None
            blocks.append(ListBlock(items=items, ordered=True, raw='\n'.join(raw_lines)))
            continue

        # ── 空行 ───────────────────────────────────────────
        if not stripped:
            blocks.append(BlankBlock(raw=line))
            pending_caption = None
            pending_code_title = None
            i += 1
            continue

        # ── 普通段落 ───────────────────────────────────────
        # 检测可能是表题（表N-N ...）或代码标题（代码清单N-N ...）
        if re.match(r'^表\s*\d+[-–]\d+', stripped) or re.match(r'^表\s*\d+[-–]\d+', stripped.lstrip('*')):
            pending_caption = stripped.strip('*').strip()
        elif re.match(r'^代码清单\s*\d+[-–]\d+', stripped):
            pending_code_title = stripped
        else:
            pending_caption = None
            pending_code_title = None

        blocks.append(ParagraphBlock(text=stripped, raw=line))
        i += 1

    return blocks


# ──────────────────────────────────────────────────────────
# 基础转换器
# ──────────────────────────────────────────────────────────

class MarkdownToDocx:
    def __init__(self, xsl_path: Optional[str] = None):
        self.doc = Document()
        self._setup_styles()
        self._load_xsl(xsl_path)
        self.stats = {'success': 0, 'failed': 0, 'total': 0}
        self.formula_errors: List[dict] = []

    def _setup_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        normal.font.size = Pt(10.5)
        normal.paragraph_format.line_spacing = 1.5

    def _load_xsl(self, xsl_path: Optional[str] = None):
        if xsl_path is None:
            xsl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MML2OMML.XSL')
        if not os.path.exists(xsl_path):
            raise FileNotFoundError(f"找不到 MML2OMML.XSL: {xsl_path}")
        xslt_doc = etree.parse(xsl_path)
        self.xslt_transform = etree.XSLT(xslt_doc)

    # ── 字体工具 ───────────────────────────────────────────

    def set_font(self, run, font_name='宋体', font_size=10.5, bold=False,
                 color=None, italic=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    # ── LaTeX → OMML ──────────────────────────────────────

    # 仅对这些"行公式"环境拆行；矩阵类环境（bmatrix/pmatrix等）保持完整
    _SPLIT_ENVS = re.compile(
        r'\\begin\{(aligned|align\*?|gather\*?|equation\*?|split)\}'
    )
    # 含有矩阵/数组类环境的公式，不做 \\ 拆行
    _MATRIX_ENVS = re.compile(
        r'\\begin\{(bmatrix|pmatrix|vmatrix|Vmatrix|matrix|array|cases|smallmatrix)\}'
    )

    def _preprocess_aligned(self, latex: str) -> List[str]:
        """
        处理多行公式环境：
        - aligned/align/gather 等：剥离环境标签，按 \\\\ 拆分为多行分别转换
        - bmatrix/pmatrix/matrix 等矩阵环境：保持完整，不拆行（\\\\ 是列分隔符）
        - 无特殊环境：直接做预处理，返回单元素列表
        """
        has_matrix = bool(self._MATRIX_ENVS.search(latex))
        has_aligned = bool(self._SPLIT_ENVS.search(latex))

        if has_matrix or not has_aligned:
            # 矩阵类或普通公式：只做常规预处理，保持 \\ 不拆分
            result = preprocess_latex(latex)
            return [result] if result else ['']

        # 行公式环境：剥离环境标签后按 \\ 拆行
        cleaned = re.sub(
            r'\\begin\{(aligned|align\*?|gather\*?|equation\*?|split)\}',
            '', latex
        )
        cleaned = re.sub(
            r'\\end\{(aligned|align\*?|gather\*?|equation\*?|split)\}',
            '', cleaned
        )
        rows = re.split(r'\\\\', cleaned)
        result = []
        for row in rows:
            row = row.replace('&', ' ').strip()
            row = preprocess_latex(row)
            if row:
                result.append(row)
        return result or ['']

    def latex_to_omml(self, latex: str):
        """LaTeX → MathML → OMML，返回可直接插入 docx 的元素，失败返回 None"""
        self.stats['total'] += 1
        latex = latex.strip()
        try:
            processed_rows = self._preprocess_aligned(latex)
            # 对每行分别转换，合并为一个 oMath
            omml_elements = []
            for row in processed_rows:
                if not row:
                    continue
                mathml_str = latex2mathml.converter.convert(row)
                mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
                omml_tree = self.xslt_transform(mathml_tree)
                omml_root = omml_tree.getroot()
                omml_xml = etree.tostring(omml_root)
                omml_element = etree.fromstring(omml_xml)
                fix_omml_runs(omml_element)
                omml_elements.append(omml_element)

            if not omml_elements:
                return None

            if len(omml_elements) == 1:
                self.stats['success'] += 1
                return omml_elements[0]

            # 多行：包裹在同一个 m:oMath 中，行间加 m:r 换行
            # 实际上 OMML 不直接支持多行，返回第一行（Word 会自动换行）
            # 对于复杂对齐公式，只返回合并的元素列表
            self.stats['success'] += 1
            return omml_elements  # 列表表示多行

        except Exception as e:
            self.stats['failed'] += 1
            preview = latex[:60]
            self.formula_errors.append({
                'latex': preview,
                'error': str(e),
            })
            print(f"  ⚠ 公式转换失败: {preview}  错误: {e}")
            return None

    # ── DOCX 渲染方法 ─────────────────────────────────────

    def render_heading(self, block: HeadingBlock):
        para = self.doc.add_heading(block.text, level=block.level)
        font_sizes = {1: 16, 2: 14, 3: 12, 4: 10.5, 5: 10.5, 6: 10.5}
        for run in para.runs:
            self.set_font(run, '黑体', font_sizes.get(block.level, 10.5), bold=True)

    def render_paragraph(self, block: ParagraphBlock):
        if not block.text.strip():
            return
        para = self.doc.add_paragraph()
        self._fill_inline(para, block.text)

    def _fill_inline(self, para, text: str):
        """在段落中填充含行内公式/代码/加粗的文本"""
        parts = re.split(r'(\$[^$]+\$|`[^`]+`|\*\*[^*]+\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('$') and part.endswith('$'):
                inner = part[1:-1].strip()
                omml = self.latex_to_omml(inner)
                if omml is not None:
                    elems = omml if isinstance(omml, list) else [omml]
                    for el in elems:
                        para._element.append(el)
                else:
                    run = para.add_run(part)
                    self.set_font(run, 'Cambria Math', 11, italic=True,
                                  color=RGBColor(255, 0, 0))
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                self.set_font(run, 'Courier New', 9)
            elif part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                self.set_font(run, '黑体', 10.5, bold=True)
            else:
                run = para.add_run(part)
                self.set_font(run)

    def render_equation(self, block: EquationBlock):
        """块级公式，居中显示"""
        omml = self.latex_to_omml(block.latex)
        elems = omml if isinstance(omml, list) else ([omml] if omml is not None else [])

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if elems:
            omathpara = OxmlElement('m:oMathPara')
            omathpara_pr = OxmlElement('m:oMathParaPr')
            jc = OxmlElement('m:jc')
            jc.set(qn('m:val'), 'center')
            omathpara_pr.append(jc)
            omathpara.append(omathpara_pr)
            for el in elems:
                omathpara.append(el)
            para._element.append(omathpara)
        else:
            run = para.add_run(f'$${block.latex}$$')
            self.set_font(run, 'Cambria Math', 11, italic=True,
                          color=RGBColor(255, 0, 0))

    def render_table(self, block: TableBlock):
        all_rows = [block.header] + block.rows
        if not block.header:
            return
        table = self.doc.add_table(rows=len(all_rows), cols=len(block.header))
        table.style = 'Table Grid'
        for i, row_data in enumerate(all_rows):
            for j, cell_text in enumerate(row_data):
                if j >= len(table.rows[i].cells):
                    continue
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_font(run, '宋体', 9)
                        if i == 0:
                            run.font.bold = True

    def render_code(self, block: CodeBlock):
        para = self.doc.add_paragraph()
        run = para.add_run(block.code)
        self.set_font(run, 'Courier New', 9)
        para.paragraph_format.left_indent = Inches(0.5)

    def render_figure(self, block: FigureBlock):
        # 尝试插入图片，失败则占位
        if block.path and os.path.exists(block.path):
            try:
                self.doc.add_picture(block.path, width=Inches(5))
            except Exception:
                self._add_figure_placeholder(block)
        else:
            self._add_figure_placeholder(block)
        # 图题（在图片下方）
        if block.caption:
            cap_para = self.doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_para.add_run(block.caption)
            self.set_font(run, '黑体', 10.5)

    def _add_figure_placeholder(self, block: FigureBlock):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = block.caption or block.alt or block.path
        run = para.add_run(f'[图片待插入: {caption}]')
        self.set_font(run, '宋体', 10.5, color=RGBColor(255, 0, 0))

    def render_list(self, block: ListBlock):
        for idx, item in enumerate(block.items):
            para = self.doc.add_paragraph()
            prefix = f'{idx + 1}. ' if block.ordered else '• '
            run = para.add_run(prefix)
            self.set_font(run)
            self._fill_inline(para, item)

    def render_blank(self, block: BlankBlock):
        pass  # 空行不添加段落

    # ── 主渲染入口 ────────────────────────────────────────

    def render_block(self, block: Block):
        if isinstance(block, HeadingBlock):
            self.render_heading(block)
        elif isinstance(block, ParagraphBlock):
            self.render_paragraph(block)
        elif isinstance(block, EquationBlock):
            self.render_equation(block)
        elif isinstance(block, TableBlock):
            self.render_table(block)
        elif isinstance(block, CodeBlock):
            self.render_code(block)
        elif isinstance(block, FigureBlock):
            self.render_figure(block)
        elif isinstance(block, ListBlock):
            self.render_list(block)
        elif isinstance(block, BlankBlock):
            self.render_blank(block)

    def convert(self, md_file: str, output_file: Optional[str] = None) -> str:
        print(f"正在转换: {md_file}")
        with open(md_file, 'r', encoding='utf-8') as f:
            md_text = f.read()

        blocks = tokenize(md_text)
        for block in blocks:
            self.render_block(block)

        if output_file is None:
            output_file = os.path.splitext(md_file)[0] + '.docx'

        self.doc.save(output_file)
        s = self.stats
        print(f"\n✅ 文档已保存: {output_file}")
        print(f"📊 公式统计: 共 {s['total']} 个, 成功 {s['success']} ✓, 失败 {s['failed']} ✗")

        if self.formula_errors:
            report_path = os.path.splitext(output_file)[0] + '_formula_report.txt'
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write("公式转换失败报告\n" + "=" * 40 + "\n\n")
                for err in self.formula_errors:
                    f.write(f"LaTeX: {err['latex']}\n错误:  {err['error']}\n\n")
            print(f"⚠  失败公式报告: {report_path}")

        return output_file


# ──────────────────────────────────────────────────────────
# 便捷函数
# ──────────────────────────────────────────────────────────

def parse_md_blocks(md_text: str) -> List[Block]:
    """将 Markdown 文本解析为 Block 列表（供外部模块使用）"""
    return tokenize(md_text)


def convert_md_to_docx(md_file: str, output_file: Optional[str] = None,
                       xsl_path: Optional[str] = None) -> str:
    converter = MarkdownToDocx(xsl_path=xsl_path)
    return converter.convert(md_file, output_file)


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Markdown → DOCX（基础转换）')
    parser.add_argument('input', help='输入 Markdown 文件')
    parser.add_argument('-o', '--output', help='输出 DOCX 文件')
    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output)
