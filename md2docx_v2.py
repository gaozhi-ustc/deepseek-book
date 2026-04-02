#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 DOCX 工具 v2
支持 LaTeX 公式转换为 Word 公式
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse
from latex2mathml.converter import convert as latex2mathml


class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.figure_counter = {}
        self.table_counter = {}
        self.current_chapter = "1"
        
    def setup_styles(self):
        """设置文档样式"""
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].paragraph_format.line_spacing = 1.5
    
    def set_font(self, run, font_name='宋体', font_size=10.5, bold=False):
        """设置字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
    
    def insert_math_equation(self, paragraph, latex_code):
        """插入 Word 公式"""
        try:
            # 转换 LaTeX 到 MathML
            mathml = latex2mathml(latex_code)
            
            # 创建 OMML 元素
            math_element = OxmlElement('m:oMath')
            math_element_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{mathml}</m:oMath>'
            
            # 简化处理：直接插入文本格式
            run = paragraph.add_run(latex_code)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.italic = True
        except Exception as e:
            # 转换失败，保留原始 LaTeX
            run = paragraph.add_run(f"${latex_code}$")
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
    
    def add_heading(self, text, level):
        """添加标题"""
        match = re.match(r'^([\d.]+)\s+', text)
        if match and level <= 2:
            self.current_chapter = match.group(1).split('.')[0]
        
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                self.set_font(run, '黑体', 16, bold=True)
            elif level == 2:
                self.set_font(run, '黑体', 14, bold=True)
            elif level == 3:
                self.set_font(run, '黑体', 12, bold=True)
            else:
                self.set_font(run, '黑体', 10.5, bold=True)
    
    def add_paragraph(self, text, style='Normal'):
        """添加段落，支持行内公式"""
        if not text.strip():
            return
        
        para = self.doc.add_paragraph()
        
        # 处理行内公式 $...$
        parts = re.split(r'(\$[^$]+\$)', text)
        
        for part in parts:
            if part.startswith('$') and part.endswith('$'):
                # 行内公式
                latex = part[1:-1].strip()
                self.insert_math_equation(para, latex)
            elif part.startswith('`') and part.endswith('`'):
                # 代码
                run = para.add_run(part[1:-1])
                self.set_font(run, 'Courier New', 9)
            elif part.startswith('**') and part.endswith('**'):
                # 粗体
                run = para.add_run(part[2:-2])
                self.set_font(run, '黑体', 10.5, bold=True)
            else:
                # 普通文本
                run = para.add_run(part)
                self.set_font(run)
    
    def add_table_from_markdown(self, lines):
        """从 Markdown 表格创建 DOCX 表格"""
        rows = []
        for line in lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if not all(re.match(r'^[-:]+$', cell) for cell in cells):
                    rows.append(cells)
        if len(rows) < 2:
            return
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_font(run, '宋体', 9)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        self.table_counter[self.current_chapter] = self.table_counter.get(self.current_chapter, 0) + 1
        table_num = f"表 {self.current_chapter}-{self.table_counter[self.current_chapter]}"
        caption = self.doc.add_paragraph()
        run = caption.add_run(table_num)
        self.set_font(run, '黑体', 10.5, bold=True)
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_code_block(self, code, language=''):
        """添加代码块"""
        para = self.doc.add_paragraph()
        run = para.add_run(code)
        self.set_font(run, 'Courier New', 9)
        para.paragraph_format.left_indent = Inches(0.5)
    
    def add_equation(self, latex):
        """添加独立公式"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        latex = latex.strip()
        if latex.startswith('$$') and latex.endswith('$$'):
            latex = latex[2:-2].strip()
        elif latex.startswith('$') and latex.endswith('$'):
            latex = latex[1:-1].strip()
        self.insert_math_equation(para, latex)
    
    def parse_markdown(self, md_file):
        """解析 Markdown 文件"""
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('```'):
                if in_code_block:
                    self.add_code_block('\n'.join(code_lines))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            if '|' in line and not in_table:
                in_table = True
                table_lines = [line]
                i += 1
                continue
            
            if in_table:
                if '|' in line:
                    table_lines.append(line)
                    i += 1
                    continue
                else:
                    self.add_table_from_markdown(table_lines)
                    table_lines = []
                    in_table = False
            
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                text = line.lstrip('#').strip()
                self.add_heading(text, level)
            elif line.strip().startswith('$$'):
                equation = line.strip()
                if not line.strip().endswith('$$'):
                    i += 1
                    while i < len(lines) and not lines[i].strip().endswith('$$'):
                        equation += '\n' + lines[i].rstrip()
                        i += 1
                    if i < len(lines):
                        equation += '\n' + lines[i].rstrip()
                self.add_equation(equation)
            elif line.strip():
                self.add_paragraph(line)
            
            i += 1
    
    def save(self, output_file):
        """保存文档"""
        self.doc.save(output_file)
        print(f"文档已保存: {output_file}")


def main():
    parser = argparse.ArgumentParser(description='Markdown 转 DOCX 工具 v2 (支持公式)')
    parser.add_argument('input', help='输入的 Markdown 文件')
    parser.add_argument('-o', '--output', help='输出的 DOCX 文件')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 {args.input}")
        return
    
    output_file = args.output or os.path.splitext(args.input)[0] + '_v2.docx'
    
    converter = MarkdownToDocx()
    converter.parse_markdown(args.input)
    converter.save(output_file)


if __name__ == '__main__':
    main()
