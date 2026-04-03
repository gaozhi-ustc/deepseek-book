#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_formatter.py — 13条规范合规层

继承 md_core.MarkdownToDocx，覆盖渲染方法以符合出版社规范：
  - 三线表（三线表）
  - 4/5级标题识别
  - 图片编号与图题（图片下方，居中）
  - 代码清单编号（代码块上方）
  - 按章编号的图/表/代码清单计数器
"""

import re
import os
from typing import Optional, List
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from md_core import (
    MarkdownToDocx, tokenize, parse_md_blocks,
    HeadingBlock, ParagraphBlock, EquationBlock,
    TableBlock, CodeBlock, FigureBlock, ListBlock, BlankBlock,
    Block,
)


# ──────────────────────────────────────────────────────────
# 三线表边框工具
# ──────────────────────────────────────────────────────────

def _make_border(val='single', sz='12', color='000000'):
    border = OxmlElement('w:top')  # 占位，调用方会替换 tag
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), sz)
    border.set(qn('w:color'), color)
    return border


def _set_table_three_line_style(table):
    """
    设置三线表样式：
      - 顶部边框 1.5pt（sz=12）
      - 底部边框 1.5pt（sz=12）
      - 表头底部 1pt（sz=8）
      - 其余边框全部取消
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 清除已有 tblBorders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')

    def make_line(tag, sz='12'):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), '000000')
        return el

    def make_none(tag):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'), 'none')
        return el

    tblBorders.append(make_line('top', '12'))
    tblBorders.append(make_none('left'))
    tblBorders.append(make_none('right'))
    tblBorders.append(make_none('insideH'))
    tblBorders.append(make_none('insideV'))
    tblBorders.append(make_line('bottom', '12'))
    tblPr.append(tblBorders)

    # 表头行底部加细线
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)


# ──────────────────────────────────────────────────────────
# 代码块灰色背景
# ──────────────────────────────────────────────────────────

def _set_gray_background(para):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)


# ──────────────────────────────────────────────────────────
# 规范合规转换器
# ──────────────────────────────────────────────────────────

class FormattedMarkdownToDocx(MarkdownToDocx):
    def __init__(self, xsl_path: Optional[str] = None):
        super().__init__(xsl_path=xsl_path)
        self.current_chapter = 0
        self.figure_counter = 0
        self.table_counter = 0
        self.code_counter = 0

    def _extract_chapter_number(self, text: str) -> Optional[int]:
        """从标题文本中提取章节号，如'第3章'→3"""
        m = re.search(r'第\s*(\d+)\s*章', text)
        if m:
            return int(m.group(1))
        # 尝试数字开头，如 "3 DeepSeek..."
        m = re.search(r'^(\d+)\s+', text)
        if m:
            return int(m.group(1))
        return None

    def _reset_counters(self):
        self.figure_counter = 0
        self.table_counter = 0
        self.code_counter = 0

    # ── 标题：识别 4/5 级别 ───────────────────────────────

    def _detect_sub_heading(self, text: str) -> Optional[int]:
        """
        检测段落文本是否是 4级 或 5级 标题。
        4级：**1. 标题** 或者直接 "1. 标题"（粗体段落）
        5级：**(1) 标题** 或者 "（1）标题"
        返回级别 4/5，或 None。
        """
        # 4级：**数字. 文字**
        if re.match(r'^\*\*\d+[.．]\s+.+\*\*$', text):
            return 4
        # 5级：**(数字) 文字** 或 **（数字）文字**
        if re.match(r'^\*\*[（(]\d+[）)]\s*.+\*\*$', text):
            return 5
        return None

    def render_heading(self, block: HeadingBlock):
        if block.level == 1:
            ch = self._extract_chapter_number(block.text)
            if ch is not None:
                self.current_chapter = ch
            else:
                self.current_chapter += 1
            self._reset_counters()

        para = self.doc.add_heading(block.text, level=block.level)
        font_sizes = {1: 16, 2: 14, 3: 12, 4: 10.5, 5: 10.5, 6: 10.5}
        for run in para.runs:
            self.set_font(run, '黑体', font_sizes.get(block.level, 10.5), bold=True)

    def render_paragraph(self, block: ParagraphBlock):
        if not block.text.strip():
            return

        # 检测是否是伪装为段落的 4/5 级标题
        level = self._detect_sub_heading(block.text)
        if level is not None:
            # 去掉 ** 包裹
            clean = re.sub(r'^\*\*|\*\*$', '', block.text).strip()
            self.render_heading(HeadingBlock(level=level, text=clean, raw=block.raw))
            return

        # 检测是否是表题、图题或代码标题行（已由上层块处理，此处跳过渲染为普通段落）
        # 但如果这些行没有被关联到对应块，仍以普通段落渲染
        para = self.doc.add_paragraph()
        self._fill_inline(para, block.text)

    # ── 三线表 ────────────────────────────────────────────

    def render_table(self, block: TableBlock):
        # 确定表题
        caption = block.caption
        if not caption:
            self.table_counter += 1
            caption = f'表{self.current_chapter}-{self.table_counter}'
        else:
            # 从现有标题提取编号
            m = re.search(r'表\s*(\d+)[-–](\d+)', caption)
            if m:
                self.table_counter = int(m.group(2))
            else:
                self.table_counter += 1

        # 表题在上方
        cap_para = self.doc.add_paragraph()
        run = cap_para.add_run(caption)
        self.set_font(run, '黑体', 10.5, bold=False)

        # 创建表格
        all_rows = [block.header] + block.rows
        if not block.header:
            return
        table = self.doc.add_table(rows=len(all_rows), cols=len(block.header))

        for i, row_data in enumerate(all_rows):
            for j, cell_text in enumerate(row_data):
                if j >= len(table.rows[i].cells):
                    continue
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_font(run, '宋体', 9)
                        if i == 0:
                            run.font.bold = True

        # 应用三线表样式
        _set_table_three_line_style(table)

    # ── 图片 ─────────────────────────────────────────────

    def render_figure(self, block: FigureBlock):
        self.figure_counter += 1

        # 确定图题
        caption = block.caption
        if not caption:
            caption = f'图{self.current_chapter}-{self.figure_counter}'
        else:
            # 若图题中已有编号，以其为准（更新计数器）
            m = re.search(r'图\s*(\d+)[-–](\d+)', caption)
            if m:
                self.figure_counter = int(m.group(2))

        # 插入图片或占位
        if block.path and os.path.exists(block.path):
            try:
                self.doc.add_picture(block.path, width=Inches(5))
            except Exception:
                self._add_figure_placeholder(block)
        else:
            self._add_figure_placeholder(block)

        # 图题在下方，居中
        cap_para = self.doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        self.set_font(run, '黑体', 10.5)

    def _add_figure_placeholder(self, block: FigureBlock):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = block.caption or block.alt or block.path
        run = para.add_run(f'[图片待插入: {label}]')
        self.set_font(run, '宋体', 10.5, color=RGBColor(180, 180, 180))

    # ── 代码清单 ──────────────────────────────────────────

    def render_code(self, block: CodeBlock):
        self.code_counter += 1

        # 确定代码标题
        title = block.title
        if not title:
            title = f'代码清单{self.current_chapter}-{self.code_counter}'
        else:
            m = re.search(r'代码清单\s*(\d+)[-–](\d+)', title)
            if m:
                self.code_counter = int(m.group(2))

        # 标题在上方
        cap_para = self.doc.add_paragraph()
        run = cap_para.add_run(title)
        self.set_font(run, '黑体', 10.5, bold=False)

        # 代码段（灰色背景，顶格，Courier New）
        para = self.doc.add_paragraph()
        run = para.add_run(block.code)
        self.set_font(run, 'Courier New', 9)
        para.paragraph_format.left_indent = Inches(0.5)
        _set_gray_background(para)


# ──────────────────────────────────────────────────────────
# 便捷函数
# ──────────────────────────────────────────────────────────

def convert_formatted(md_file: str, output_file: Optional[str] = None,
                      xsl_path: Optional[str] = None) -> str:
    converter = FormattedMarkdownToDocx(xsl_path=xsl_path)
    return converter.convert(md_file, output_file)


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Markdown → DOCX（符合13条规范）')
    parser.add_argument('input', help='输入 Markdown 文件')
    parser.add_argument('-o', '--output', help='输出 DOCX 文件')
    args = parser.parse_args()
    convert_formatted(args.input, args.output)
